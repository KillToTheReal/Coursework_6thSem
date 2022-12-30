import gzip
import pandas as pd
import pypdfium2 as pdfium
from lxml import etree
import time
import os
import docx
import codecs
from search.documents import WikiPage, TextFile, PDFFile, XLSFile, DOCXFile
import openpyxl
def load_documents(path_from):
    start_path = path_from # current directory
    start = time.time()
    doc_id = 0
    for path,dirs,files in os.walk(start_path):
        for filename in files:
            path_to_file = os.path.join(path,filename)
            
            if filename.endswith('.txt'):
                with open(path_to_file) as f:
                    abstract = f.read()
                    yield TextFile(ID=doc_id, abstract=abstract, path=path_to_file)
                    doc_id += 1
                    
            elif filename.endswith('.pdf'):
                pdf = pdfium.PdfDocument(path_to_file)
                pdf_page = 1
                for i in range(len(pdf)):
                    abstract = pdf[i].get_textpage().get_text_range()
                    yield PDFFile(ID=doc_id, abstract=abstract, page=pdf_page, path=path_to_file)
                    pdf_page += 1 
                    doc_id += 1
                    
            elif not filename.startswith('~') and filename.endswith('.docx'):
                doc = docx.Document(path_to_file)
                parag = 1
                for i in doc.paragraphs:
                    abstract = i.text
                    yield DOCXFile(ID=doc_id, abstract=abstract, paragraph=parag, path=path_to_file)
                    parag += 1 
                    doc_id += 1    

            # elif filename == 'enwiki-latest-abstract.xml.gz':
            #     with gzip.open('data/enwiki-latest-abstract.xml.gz', 'rb') as f:   
            #         for _, element in etree.iterparse(f, events=('end',), tag='doc'):
            #             title = element.findtext('./title')
            #             url = element.findtext('./url')
            #             abstract = element.findtext('./abstract')
            #             path = './data/123.xml'
            #             yield WikiPage(ID=doc_id, title=title, url=url, abstract=abstract, path=path)
            #             doc_id += 1
            #             element.clear()
    end = time.time()
    print(f'Parsing Documents took {end - start} seconds')
